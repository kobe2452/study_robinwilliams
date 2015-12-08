import docx, re

doc_serial = re.compile(r'\d+ of \d+ DOCUMENTS', re.M|re.I)
date = re.compile(r'\w+ \d+, \d\d\d\d \w+', re.M|re.I)

# pattern = re.compile(r'^[A-Z\-]+:', re.M|re.I)
pattern = re.compile(r'^[A-Z]+\-?[A-Z]+:', re.M|re.I)

filename = '/Users/tl8313/Documents/study_robinwilliams/doc/US_Newspapers_Magazine_Stories.docx'

document = docx.Document(filename)
# print type(document.paragraphs), len(document.paragraphs)
for index, paragraph in enumerate(document.paragraphs):
    para = paragraph.text.encode('utf-8')

    # if pattern.match(para):
    #     print para

    if doc_serial.search(para):
        serial = para.split()[0]
    if date.match(para):
        date_nextline = document.paragraphs[index+1].text.encode('utf-8')
    if para.startswith('BYLINE'):
        byline = para
        title_line = document.paragraphs[index-2].text.encode('utf-8')
        print serial, title_line
        print
    if para.startswith('SECTION'):
        section = para
    if para.startswith('LENGTH'):
        length = para
    if para.startswith('DATELINE'):
        dateline = para
    if para.startswith('LANGUAGE'):
        language = para
    if para.startswith('PUBLICATION-TYPE'):
        publication = para
    if para.startswith('SUBJECT'):
        subject = para
    if para.startswith('STATE'):
        state = para
    if para.startswith('COUNTRY'):
        country = para
    if para.startswith('LOAD-DATE'):
        loaddate = para


# docText = '\n'.join([
#     paragraph.text.encode('utf-8') for paragraph in document.paragraphs
# ])

# print docText